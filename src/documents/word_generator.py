"""Generate Hebrew Word reports using docxtpl.

The template (data/templates/סיכום בדיקת התכנות טמפלט.docx) already has
full Hebrew RTL formatting. We fill in Jinja2 tags with report data.

All string values use UTF-8 encoding.  RTL handling is entirely via the
template -- this module does NOT set paragraph direction programmatically.
"""

from __future__ import annotations

import logging
import os
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING, Any

from docxtpl import DocxTemplate

if TYPE_CHECKING:
    from src.models.report import AuditEntry, ReportData

logger = logging.getLogger(__name__)


class WordGenerator:
    """Generate Hebrew Word reports using docxtpl.

    The Word template already contains full Hebrew RTL formatting.
    We just populate the Jinja2 context variables.
    """

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def generate_report(
        self,
        report_data: ReportData,
        template_path: str,
        output_path: str,
    ) -> str:
        """Generate the main feasibility report.

        Maps ReportData fields to template context variables and renders
        the Word document.

        Args:
            report_data: Complete report data.
            template_path: Path to the .docx template with Jinja2 tags.
            output_path: Destination path for the generated document.

        Returns:
            The output file path.
        """
        self._validate_template(template_path)
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

        context = self._build_context(report_data)

        tpl = DocxTemplate(template_path)
        tpl.render(context)
        tpl.save(output_path)

        logger.info("Report generated: %s", output_path)
        return output_path

    def generate_audit_log_doc(
        self,
        audit_entries: list[AuditEntry],
        report_data: ReportData,
        output_path: str,
    ) -> str:
        """Generate a companion audit log document.

        Creates a simple Word document listing every audit entry
        (calculation tool invocations, their inputs, formulas, and results).

        Args:
            audit_entries: List of audit entries.
            report_data: The report data (for header info).
            output_path: Destination path.

        Returns:
            The output file path.
        """
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

        from docx import Document

        doc = Document()

        # Title — Hebrew labels for professional output
        doc.add_heading("יומן חישובים", level=1)
        doc.add_paragraph(f"תאריך דוח: {self._format_date(report_data.report_date)}")
        doc.add_paragraph(f"מספר רשומות: {len(audit_entries)}")
        doc.add_paragraph("")

        for idx, entry in enumerate(audit_entries, 1):
            doc.add_heading(f"רשומה {idx}: {entry.tool_name}", level=2)
            doc.add_paragraph(f"חותמת זמן: {entry.timestamp}")
            doc.add_paragraph(f"נוסחה: {entry.formula}")
            doc.add_paragraph(f"קלטים: {self._dict_to_str(entry.inputs)}")
            doc.add_paragraph(f"שיעורים בשימוש: {self._dict_to_str(entry.rates_used)}")
            doc.add_paragraph(f"תוצאה: {self._dict_to_str(entry.result)}")
            if entry.user_overrides:
                doc.add_paragraph(f"שינויי משתמש: {self._dict_to_str(entry.user_overrides)}")
            if entry.reasoning:
                doc.add_paragraph(f"הנמקה: {entry.reasoning}")
            if entry.source_reference:
                doc.add_paragraph(f"מקור: {entry.source_reference}")
            doc.add_paragraph("")

        doc.save(output_path)
        logger.info("Audit log document generated: %s", output_path)
        return output_path

    # ------------------------------------------------------------------
    # Context building
    # ------------------------------------------------------------------

    def _build_context(self, report_data: ReportData) -> dict[str, Any]:
        """Build the template context dict from ReportData.

        Maps model fields to template variable names.
        Formats numbers with comma separators and dates as DD/MM/YYYY.
        """
        nachla = report_data.nachla
        nachla_dict: dict[str, Any] = {}
        if hasattr(nachla, "model_dump"):
            nachla_dict = nachla.model_dump()
        elif isinstance(nachla, dict):
            nachla_dict = nachla

        # Basic info
        context: dict[str, Any] = {
            # Header
            "owner_name": nachla_dict.get("owner_name", ""),
            "moshav_name": nachla_dict.get("moshav_name", ""),
            "gush": nachla_dict.get("gush", ""),
            "helka": nachla_dict.get("helka", ""),
            "report_date": self._format_date(report_data.report_date),
            "report_date_iso": report_data.report_date,
            # Study objectives (section 2)
            "study_objectives": report_data.study_objectives,
            # Disclaimers (section 3) — includes both header and mandatory disclaimers
            "disclaimers": report_data.format_disclaimers(
                report_date=self._format_date(report_data.report_date),
                moshav_name=nachla_dict.get("moshav_name", ""),
                data_year=str(datetime.now().year),
                survey_date=self._format_date(nachla_dict.get("survey_map_date", "")),
            ),
            "header_disclaimers": self._format_header_disclaimers(
                report_data, nachla_dict
            ),
            # Taba analysis (section 4)
            "tabas": self._format_tabas(report_data),
            # Capitalization (section 5)
            "hivun_375_result": report_data.hivun_375_result or {},
            "hivun_33_result": report_data.hivun_33_result or {},
            "hivun_comparison": report_data.hivun_comparison or {},
            # Buildings (section 6-7)
            "buildings": self._format_buildings(report_data),
            "building_cards": self._format_building_cards(report_data),
            # Cost summary (section 8)
            "total_regularization_cost": self._format_currency(report_data.total_regularization_cost),
            "total_usage_fees": self._format_currency(report_data.total_usage_fees),
            "total_permit_fees": self._format_currency(report_data.total_permit_fees),
            "total_regularization_cost_raw": report_data.total_regularization_cost,
            "total_usage_fees_raw": report_data.total_usage_fees,
            "total_permit_fees_raw": report_data.total_permit_fees,
            # Split (section 9)
            "split_results": report_data.split_results,
            # Recommendations (section 10)
            "recommendations": report_data.recommendations,
            "action_items": self._format_action_items(report_data),
            # Appendices (section 11)
            "appendix_building_sizes": report_data.appendix_building_sizes,
            "annotated_survey_map_path": report_data.annotated_survey_map_path or "",
            # Warnings
            "missing_data_warnings": report_data.missing_data_warnings,
            # Legal / authorization
            "authorization_type": nachla_dict.get("authorization_type", ""),
            "is_capitalized": nachla_dict.get("is_capitalized", False),
            "capitalization_track": nachla_dict.get("capitalization_track", ""),
            "priority_area": nachla_dict.get("priority_area", "none"),
            "num_existing_houses": nachla_dict.get("num_existing_houses", 0),
        }
        return context

    # ------------------------------------------------------------------
    # Formatting helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _format_currency(amount: float) -> str:
        """Format as Hebrew currency: '1,234,567 \\u20aa'."""
        if amount == 0:
            return "0 \u20aa"
        return f"{amount:,.0f} \u20aa"

    @staticmethod
    def _format_date(iso_date: str) -> str:
        """Convert ISO date (YYYY-MM-DD) to DD/MM/YYYY format."""
        if not iso_date:
            return ""
        try:
            dt = datetime.fromisoformat(iso_date)
            return dt.strftime("%d/%m/%Y")
        except (ValueError, TypeError):
            return iso_date

    def _format_buildings(self, report_data: ReportData) -> list[dict[str, Any]]:
        """Format building data for template rendering."""
        result: list[dict[str, Any]] = []
        for b in report_data.buildings:
            bd: dict[str, Any] = {}
            if hasattr(b, "model_dump"):
                bd = b.model_dump()
            elif isinstance(b, dict):
                bd = b
            result.append(bd)
        return result

    def _format_building_cards(self, report_data: ReportData) -> list[dict[str, Any]]:
        """Format building cards for template rendering."""
        result: list[dict[str, Any]] = []
        for card in report_data.building_cards:
            cd: dict[str, Any] = {}
            if hasattr(card, "model_dump"):
                cd = card.model_dump()
            elif isinstance(card, dict):
                cd = card
            # Format monetary values
            for key in ("permit_fees", "usage_fees", "betterment_levy", "total_cost"):
                if key in cd and isinstance(cd[key], (int, float)):
                    cd[f"{key}_formatted"] = self._format_currency(cd[key])
            result.append(cd)
        return result

    def _format_action_items(self, report_data: ReportData) -> list[dict[str, Any]]:
        """Format action items for template rendering."""
        result: list[dict[str, Any]] = []
        for item in report_data.action_items:
            ad: dict[str, Any] = {}
            if hasattr(item, "model_dump"):
                ad = item.model_dump()
            elif isinstance(item, dict):
                ad = item
            result.append(ad)
        return result

    def _format_tabas(self, report_data: ReportData) -> list[dict[str, Any]]:
        """Format taba/zoning plan data for template rendering (section 4)."""
        result: list[dict[str, Any]] = []
        for t in report_data.tabas:
            td: dict[str, Any] = {}
            if hasattr(t, "model_dump"):
                td = t.model_dump()
            elif isinstance(t, dict):
                td = t
            result.append(td)
        return result

    def _format_header_disclaimers(
        self, report_data: ReportData, nachla_dict: dict[str, Any]
    ) -> list[str]:
        """Format the section-3 header disclaimers from REPORT_HEADER_DISCLAIMERS."""
        from src.models.report import REPORT_HEADER_DISCLAIMERS

        formatted: list[str] = []
        for disclaimer in REPORT_HEADER_DISCLAIMERS:
            try:
                formatted.append(disclaimer.format(
                    moshav_name=nachla_dict.get("moshav_name", ""),
                    data_year=str(datetime.now().year),
                    survey_date=self._format_date(nachla_dict.get("survey_map_date", "")),
                ))
            except KeyError:
                formatted.append(disclaimer)
        return formatted

    @staticmethod
    def _dict_to_str(d: dict[str, Any]) -> str:
        """Convert a dict to a readable string for audit log documents."""
        import json

        return json.dumps(d, ensure_ascii=False, default=str, indent=2)

    # ------------------------------------------------------------------
    # Validation
    # ------------------------------------------------------------------

    @staticmethod
    def _validate_template(template_path: str) -> None:
        """Validate that the template file exists and is a .docx."""
        path = Path(template_path)
        if not path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")
        if path.suffix.lower() != ".docx":
            raise ValueError(f"Template is not a .docx file: {template_path}")
