"""Build professional Hebrew Word reports from calculation data.

Generates a .docx file matching the professional template style with
all sections filled from actual process data. Uses python-docx directly
since the template has no Jinja2 tags.
"""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def _set_rtl(paragraph):
    """Set paragraph to RTL direction."""
    pPr = paragraph._element.get_or_add_pPr()
    bidi = pPr.makeelement(qn('w:bidi'), {})
    pPr.append(bidi)


def _add_rtl_paragraph(doc, text="", style=None, bold=False, size=None, color=None, alignment=None):
    """Add an RTL paragraph with optional formatting."""
    p = doc.add_paragraph(text, style=style)
    _set_rtl(p)
    if alignment:
        p.alignment = alignment
    if text and (bold or size or color):
        run = p.runs[0] if p.runs else p.add_run(text)
        if bold:
            run.bold = True
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p


def _fmt_currency(amount: float) -> str:
    """Format as Israeli currency."""
    if amount == 0:
        return "0 \u20aa"
    return f"{amount:,.0f} \u20aa"


def _fmt_date(iso_date: str) -> str:
    """Convert ISO date to DD/MM/YYYY."""
    if not iso_date:
        return ""
    try:
        return datetime.fromisoformat(iso_date).strftime("%d/%m/%Y")
    except (ValueError, TypeError):
        return iso_date


def build_report(
    nachla: dict[str, Any],
    buildings: list[dict[str, Any]],
    tabas: list[dict[str, Any]],
    calc_results: dict[str, Any],
    report_date: str,
    output_path: str,
) -> str:
    """Build a professional Hebrew feasibility report.

    Args:
        nachla: Nachla data dict (owner_name, moshav_name, etc.)
        buildings: List of building dicts
        tabas: List of taba dicts
        calc_results: All calculation results
        report_date: Report date string
        output_path: Output .docx path

    Returns:
        The output file path.
    """
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'David'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:cs'), 'David')

    owner = nachla.get('owner_name', '')
    moshav = nachla.get('moshav_name', '')
    gush = nachla.get('gush', '')
    helka = nachla.get('helka', '')
    priority = nachla.get('priority_area', 'none')

    # ── Title ──
    title = _add_rtl_paragraph(doc, f"בדיקת התכנות משפחת {owner} מושב {moshav}",
                                bold=True, size=18, color=(59, 89, 39))
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_rtl_paragraph(doc, f"גוש {gush} חלקה {helka} | תאריך: {_fmt_date(report_date)}",
                       size=10, color=(120, 120, 120))
    doc.add_paragraph()

    # ── Section 1: Study Objectives ──
    _add_rtl_paragraph(doc, "מטרות בדיקת ההתכנות:", bold=True, size=13)
    goals = nachla.get('client_goals', [])
    goal_labels = {
        'regularization': 'יצירת מתווה הסדרה למשק והערכת עלויות',
        'capitalization': 'תחשיבים להיוון 3.75% + היוון 33%',
        'split': 'מתווה לפיצול מגרשים בנחלה',
        'all': 'בדיקה מקיפה של כלל האפשרויות',
    }
    for g in goals:
        _add_rtl_paragraph(doc, f"• {goal_labels.get(g, g)}")
    doc.add_paragraph()

    # ── Section 2: Disclaimers ──
    _add_rtl_paragraph(doc, "הצהרות ומגבלות:", bold=True, size=13)
    disclaimers = [
        f"שווי הזכויות בנחלה הסתמך על נתונים במושב {moshav}. בדיקת ההתכנות בוצעה על סמך גבולות התב\"ע הנוכחית.",
        "מזמין העבודה מודע לכך שלא בוצעה שומה למשק, ושהערכות אלו בוצעו לשם קבלת סדר גודל ראשוני בלבד ואינם מהוות אסמכתא משפטית.",
        "בדיקת התכנות זו הוכנה עבור מזמינה ולמטרתה בלבד. אין היא מהווה תחליף לייעוץ משפטי ו/או לשומה.",
        f"תוקף הבדיקה: 6 חודשים מיום {_fmt_date(report_date)}.",
    ]
    for d in disclaimers:
        p = _add_rtl_paragraph(doc, d, size=9, color=(100, 100, 100))
    doc.add_paragraph()

    # ── Section 3: Taba Analysis ──
    _add_rtl_paragraph(doc, 'ניתוח תב"ע קיימת:', bold=True, size=13)
    if tabas:
        for t in tabas:
            taba_name = t.get('taba_name', '')
            taba_num = t.get('taba_number', '')
            _add_rtl_paragraph(doc, f'תב"ע {taba_num} — {taba_name}', bold=True, size=11)

            plot = t.get('plot_size_sqm', 0)
            units = t.get('num_units_allowed', 0)
            _add_rtl_paragraph(doc, f"שטח מגרש: {plot:,.0f} מ\"ר | יחידות דיור מותרות: {units}")

            unit_rights = t.get('unit_rights', [])
            if unit_rights:
                ur = unit_rights[0] if isinstance(unit_rights[0], dict) else {}
                main = ur.get('main_area_sqm', 0)
                service = ur.get('service_area_sqm', 0)
                _add_rtl_paragraph(doc, f"זכויות ליח' דיור: {main} מ\"ר עיקרי + {service} מ\"ר שירות")

            # Enhancement 6: attached_unit_allowed flag
            flags = []
            if t.get('split_allowed'):
                flags.append("פיצול מותר")
            if t.get('pool_allowed'):
                flags.append("בריכה מותרת")
            if t.get('attached_unit_allowed'):
                flags.append("יחידה צמודה מותרת")
            if flags:
                _add_rtl_paragraph(doc, "אפשרויות: " + " | ".join(flags))
        doc.add_paragraph()
    else:
        _add_rtl_paragraph(doc, 'לא הוזנו נתוני תב"ע.', color=(180, 100, 0))
        doc.add_paragraph()

    # ── Section 4: Building Status ──
    _add_rtl_paragraph(doc, "סטטוס מבנים בנחלה:", bold=True, size=13)
    if buildings:
        # Enhancement 1: Expanded 8-column building table
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        headers = ['#', 'שם מבנה', 'סוג', 'שטח עיקרי', 'שטח שירות', 'ממ"ד', 'סטטוס', 'סטייה']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                _set_rtl(p)
                p.runs[0].bold = True

        type_labels = {
            'residential': 'בית מגורים', 'service': 'שירות/מחסן',
            'agricultural': 'חקלאי', 'pool': 'בריכה', 'pergola': 'פרגולה',
            'plach': 'פל"ח', 'shed_open': 'סככה פתוחה',
        }
        status_labels = {
            'compliant': 'תקין', 'deviation': 'חריגה',
            'no_permit': 'ללא היתר', 'marked_demolition': 'להריסה',
        }
        for b in buildings:
            row = table.add_row()
            row.cells[0].text = str(b.get('id', ''))
            row.cells[1].text = b.get('name', '')
            row.cells[2].text = type_labels.get(b.get('building_type', ''), b.get('building_type', ''))
            row.cells[3].text = f"{b.get('main_area_sqm', 0)} מ\"ר"
            row.cells[4].text = f"{b.get('service_area_sqm', 0)} מ\"ר"
            row.cells[5].text = f"{b.get('mamad_area_sqm', 0)} מ\"ר"
            row.cells[6].text = status_labels.get(b.get('status', ''), b.get('status', ''))
            deviation = b.get('deviation_sqm', 0)
            row.cells[7].text = f"{deviation} מ\"ר" if deviation else "-"
            for cell in row.cells:
                for p in cell.paragraphs:
                    _set_rtl(p)
        doc.add_paragraph()
    else:
        _add_rtl_paragraph(doc, "לא זוהו מבנים.", color=(180, 100, 0))
        doc.add_paragraph()

    # ── Section 4b: Split (Enhancement 2) ──
    split_data = calc_results.get('split', {})
    if split_data:
        _add_rtl_paragraph(doc, "פיצול נחלה:", bold=True, size=13)
        eligibility = split_data.get('eligibility', {})
        is_eligible = eligibility.get('eligible', False)
        reason = eligibility.get('reason', '')

        if is_eligible:
            _add_rtl_paragraph(doc, "הנחלה זכאית לפיצול.")
            split_cost = split_data.get('cost', {}).get('result', 0)
            if split_cost:
                _add_rtl_paragraph(doc, f"עלות פיצול מוערכת: {_fmt_currency(split_cost)}")
        else:
            _add_rtl_paragraph(doc, f"הנחלה אינה זכאית לפיצול: {reason}")

        # bar_reshut warning
        auth_type = nachla.get('authorization_type', '')
        client_goals = nachla.get('client_goals', [])
        if auth_type == 'bar_reshut' and 'split' in client_goals:
            warn_p = _add_rtl_paragraph(doc, "בר-רשות אינו רשאי לפצל ללא חוזה חכירה — נדרש היוון 33% תחילה",
                                        bold=True, color=(255, 0, 0))
        doc.add_paragraph()

    # ── Section 5: Context Notes (Enhancement 5) ──
    _add_rtl_paragraph(doc, "הערות כלליות:", bold=True, size=13)

    is_capitalized = nachla.get('is_capitalized', False)
    if is_capitalized:
        cap_track = nachla.get('capitalization_track', '')
        _add_rtl_paragraph(doc, f"הנחלה מהוונת במסלול {cap_track}" if cap_track else "הנחלה מהוונת")
    else:
        _add_rtl_paragraph(doc, "הנחלה אינה מהוונת כיום")

    num_houses = nachla.get('num_existing_houses', 0)
    _add_rtl_paragraph(doc, f"{num_houses} בתי מגורים קיימים")

    _add_rtl_paragraph(doc, 'ממ"ד ראשון בכל בית — פטור עד 12 מ"ר מדמי היתר')
    doc.add_paragraph()

    # ── Section 6: Hivun (Capitalization) ──
    _add_rtl_paragraph(doc, "היוון המשק:", bold=True, size=13)

    hivun = calc_results.get('hivun', {})
    h375 = hivun.get('hivun_375', {})
    h33 = hivun.get('hivun_33', {})
    h375_total = h375.get('result', 0) if isinstance(h375, dict) else 0
    h33_total = h33.get('result', 0) if isinstance(h33, dict) else 0

    _add_rtl_paragraph(doc, "מסלול דמי חכירה — 3.75%:", bold=True, size=11)
    _add_rtl_paragraph(doc, "במסלול זה בעל הנחלה רוכש את זכויות הבנייה הבסיסיות.")
    _add_rtl_paragraph(doc, f"עלות מוערכת: {_fmt_currency(h375_total)}", bold=True)
    doc.add_paragraph()

    _add_rtl_paragraph(doc, "מסלול דמי רכישה — 33%:", bold=True, size=11)
    _add_rtl_paragraph(doc, "במסלול זה בעל הנחלה רוכש את מלוא הזכויות כולל אפשרות לפיצול.")
    _add_rtl_paragraph(doc, f"עלות מוערכת: {_fmt_currency(h33_total)}", bold=True)
    doc.add_paragraph()

    if h375_total > 0 and h33_total > 0:
        diff = h33_total - h375_total
        _add_rtl_paragraph(doc, f"הפרש בין המסלולים: {_fmt_currency(diff)}")
        if h375_total < h33_total:
            _add_rtl_paragraph(doc, "מסלול 3.75% זול יותר בכניסה, אך פיצול ידרוש תשלום נוסף.")
        doc.add_paragraph()

    # ── Section 7: Cost Summary (Enhancement 3) ──
    _add_rtl_paragraph(doc, "סיכום עלויות:", bold=True, size=13)

    usage = calc_results.get('usage_fees', {}).get('total', 0)
    permit = calc_results.get('regularization', {}).get('total_permit_fees', 0)
    # Only include split cost if eligible
    split_eligible = calc_results.get('split', {}).get('eligibility', {}).get('eligible', False)
    split_cost = calc_results.get('split', {}).get('cost', {}).get('result', 0) if split_eligible else 0

    # Sum all betterment results
    betterment_data = calc_results.get('betterment', {})
    betterment_total = sum(
        v.get('result', 0) for v in betterment_data.values() if isinstance(v, dict)
    )

    # Hivun: use the lower of the two tracks
    hivun_lower = min(h375_total, h33_total) if h375_total > 0 and h33_total > 0 else max(h375_total, h33_total)

    grand_total = usage + permit + betterment_total + split_cost + hivun_lower

    cost_table = doc.add_table(rows=1, cols=2)
    cost_table.style = 'Table Grid'
    cost_table.rows[0].cells[0].text = "סוג תשלום"
    cost_table.rows[0].cells[1].text = 'סכום (ש"ח)'
    for p in cost_table.rows[0].cells[0].paragraphs:
        _set_rtl(p)
        p.runs[0].bold = True
    for p in cost_table.rows[0].cells[1].paragraphs:
        _set_rtl(p)
        p.runs[0].bold = True

    costs = [
        ("דמי שימוש", usage),
        ("דמי היתר", permit),
        ("היוון 3.75%", h375_total),
        ("היוון 33%", h33_total),
        ("פיצול", split_cost),
        ("היטל השבחה", betterment_total),
    ]
    for label, amount in costs:
        row = cost_table.add_row()
        row.cells[0].text = label
        row.cells[1].text = _fmt_currency(amount)
        for cell in row.cells:
            for p in cell.paragraphs:
                _set_rtl(p)

    # Grand total row (bold)
    total_row = cost_table.add_row()
    total_row.cells[0].text = "סה\"כ מוערך"
    total_row.cells[1].text = _fmt_currency(grand_total)
    for cell in total_row.cells:
        for p in cell.paragraphs:
            _set_rtl(p)
            for run in p.runs:
                run.bold = True

    _add_rtl_paragraph(doc, "* מסלולי ההיוון (3.75% ו-33%) הם חלופיים — הסה\"כ כולל את הנמוך מביניהם.",
                       size=9, color=(100, 100, 100))
    doc.add_paragraph()

    # ── Section 8: Priority Area (Enhancement 4) ──
    if priority and priority != 'none':
        area_labels = {'A': 'א', 'B': 'ב', 'frontline': 'קו עימות'}
        _add_rtl_paragraph(doc, "הנחות אזור עדיפות לאומית:", bold=True, size=13)
        _add_rtl_paragraph(doc, f"מושב {moshav} מוגדר תחת אזור עדיפות לאומית {area_labels.get(priority, priority)}.")
        _add_rtl_paragraph(doc, "בישובים המוגדרים כאזורי עדיפות לאומית ישנן הטבות במספר תשלומים עבור רמ\"י.")

        # Quantified priority discounts
        priority_discount_375 = h375.get('priority_discount_applied', 0) if isinstance(h375, dict) else 0
        rate_applied_33 = h33.get('rate_applied', 0) if isinstance(h33, dict) else 0

        if priority_discount_375:
            _add_rtl_paragraph(doc, f"הנחת היוון 3.75%: {priority_discount_375 * 100:.0f}%")
        if rate_applied_33:
            _add_rtl_paragraph(doc, f"שיעור דמי רכישה: {rate_applied_33 * 100:.2f}% במקום 33%")
        doc.add_paragraph()

    # ── Section 9: Footer ──
    _add_rtl_paragraph(doc, "───────────────────────────────")
    p = _add_rtl_paragraph(doc, f"דוח זה הופק באמצעות מערכת בדיקת התכנות נחלות | {_fmt_date(report_date)}",
                           size=8, color=(150, 150, 150))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    logger.info("Professional report built: %s", output_path)
    return output_path
